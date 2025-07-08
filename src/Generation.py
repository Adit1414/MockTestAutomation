from openai import OpenAI
import random
import time
import textwrap
import os
import pandas as pd
from docx import Document
from PromptsDict import prompt_templates  # Load the prompt dictionary
import winsound
from datetime import datetime

# === CONFIGURATION ===
MODEL = 'gpt-4-turbo'
EXCEL_PATH = "CS/Gen/Syllabus.xlsx"
CHUNK_SIZE = 5
TESTING = True
EXAM = "UGC NET"
QUESTIONS_PATH = "CS/Gen/Questions.docx"
VERIFICATIONS_PATH = "CS/Gen/Verifications.docx"
SKIPPED_PATH = "CS/Gen/Skipped.docx"
skipped_chunks = []


# === UTILITIES ===
def get_user_question_plan():
    print("📋 Available Question Types:")
    for i, qtype in enumerate(prompt_templates.keys(), start=1):
        print(f"{i}. {qtype}")
    
    plan = {}
    while True:
        try:
            qtype_input = input("✍️ Enter question type name (or press Enter to finish): ").strip()
            if qtype_input == "":
                break
            if qtype_input not in prompt_templates:
                print("❌ Invalid question type. Try again.")
                continue
            count = int(input(f"🔢 How many questions of type '{qtype_input}' do you want? (must be multiple of {CHUNK_SIZE}): "))
            if count % CHUNK_SIZE != 0:
                print(f"⚠️ Must be a multiple of {CHUNK_SIZE}. Try again.")
                continue
            plan[qtype_input] = plan.get(qtype_input, 0) + count
        except ValueError:
            print("❌ Invalid number. Please try again.")
    return plan

def load_all_topics(excel_path=EXCEL_PATH):
    try:
        df = pd.read_excel(excel_path)
        topic_column = df.columns[0]
        topics = df[topic_column].dropna().tolist()
        random.shuffle(topics)
        return topics
    except Exception as e:
        print(f"❌ Error loading topics: {e}")
        return []

def validate_topic_capacity(plan, total_topics):
    total_chunks_requested = sum(num // CHUNK_SIZE for num in plan.values())
    if total_chunks_requested > len(total_topics) // CHUNK_SIZE:
        print(f"❌ Not enough unique topics to generate all question chunks without duplication.")
        print(f"💡 Topics available: {len(total_topics)}, Questions requested: {total_chunks_requested * CHUNK_SIZE}")
        winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
        exit(1)

def format_topics(topics_list):
    return "\n".join([f"{i+1}. {topic}" for i, topic in enumerate(topics_list)]) 

def build_prompt_from_template(topics_list, template_key, num_of_questions):
    topics_str = format_topics(topics_list)
    randomized_answer_key = ', '.join(str(n) for n in random.choices(range(1, 5), k=CHUNK_SIZE))
    template = prompt_templates.get(template_key, "")
    return template.format(topics=topics_str, answer_key=randomized_answer_key, num = num_of_questions, exam = EXAM)

def generate_all_prompts_from_plan(plan, all_topics):
    all_prompts = []
    topic_index = 0

    for qtype, total_questions in plan.items():
        num_chunks = total_questions // CHUNK_SIZE
        print(f"\n🧠 Generating {total_questions} '{qtype}' questions in {num_chunks} chunks...")
        for _ in range(num_chunks):
            topics_chunk = all_topics[topic_index: topic_index + CHUNK_SIZE]
            topic_index += CHUNK_SIZE
            prompt = build_prompt_from_template(topics_chunk, qtype, "five")
            all_prompts.append((qtype, prompt))
    return all_prompts

def check_file_access(filename):
    try:
        with open(filename, 'a'):
            pass
    except Exception as e:
        raise IOError(f"❌ Cannot access {filename}. Please close it if open. Details: {e}")

def call_gpt(prompt, testing, chunks, retries=3):
    if testing:
        print("⚠️ MOCK GPT CALLED.")
        if "VERIFICATION" in prompt:
            return "\n".join([
                f"Question {i+1} - Correct" for i in range(chunks)
            ])
        else:
            return "\n\n".join(
                [f"--Question Starting--\nQ{i+1}. Sample question on Topic.\nAnswer: {chr(65 + i%4)}\nExplanation: Because XYZ."
                 for i in range(chunks)]
            )
    else:
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        for attempt in range(retries):
            try:
                response = client.chat.completions.create(
                    model=MODEL,
                    messages=[
                        {"role": "system", "content": "You are a UGC NET paper setter."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=4000
                )
                return response.choices[0].message.content
            except Exception as e:
                print(f"⚠️ GPT attempt {attempt+1} failed: {e}")
                time.sleep(2)
        raise RuntimeError("❌ All GPT retries failed. Aborting to prevent wasted usage.")
   
def verify_prompt(questions_text):
    return f"""
VERIFICATION: im gonna send you 5 questions at once
you have to check whether or not they are correct
question, answer key, solution etc all should be 100% correct, if there is even a tiny fault or discrepancy, tell me
be very precise and unbiased, check very carefully
you dont need to repeat the questions and answers
just thoroughly check them if they are written correctly and then tell me for each question if it is correct or not, free of discrepancies or not. Are the answer keys correct? Are the solutions logical? Are any options misleading or repeated?
if there are any issues, tell me what and where they are, then tell me how to fix them
Example format for verifying
Question 1 - Correct
Question 2 - Correct
Question 3 - INCORRECT - ANSWER KEY MISMATCH //or whatever the problem was
Question 4 - Correct
Question 5 - Correct
<Rewritten Version of the incorrect question - whole question, answer and solution of any incorrect questions>
so final format for each question:
Question <Q number> - <Correctness> - <Problem if incorrect>

this is only an example, be completely unbiased in determining the correctness and incorrectness. check and recheck multiple times the correctness
Here are the questions:
{questions_text}
"""
   
def save_response(text, folder="CS/RawResponses"):
    # Make sure folder exists
    os.makedirs(folder, exist_ok=True)
    
    # Create timestamped filename
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"gpt_response_{timestamp}.docx"
    
    # Full path to save the file
    path = os.path.join(folder, filename)

    # Save to Word file
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    print(f"✅ Response saved to: {path}")
             
def save_to_docx(content, filename="Questions.docx"):
    try:
        with open(filename, 'a'):
            pass  # Attempt to open for append to check if it's locked
    except Exception as e:
        raise IOError(f"❌ Cannot access {filename}. Make sure it is closed. Details: {e}")

    try:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filename)
    except Exception as e:
        raise IOError(f"❌ Failed to save {filename}. Details: {e}")

# === MAIN EXECUTION ===
def GenerationMain(chunks):
    user_plan = get_user_question_plan()
    all_topics = load_all_topics()
    validate_topic_capacity(user_plan, all_topics)
    generated_prompts = generate_all_prompts_from_plan(user_plan, all_topics)

    # ✅ Check access to Word files BEFORE any GPT calls
    try:
        check_file_access(QUESTIONS_PATH)
        check_file_access(VERIFICATIONS_PATH)
    except Exception as e:
        print(e)
        winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
        exit(1)
    all_questions = []

    for qtype, prompt in generated_prompts:
        print(f"{qtype} : {prompt} \n\n")
        try:
            generated_chunk = call_gpt(prompt, TESTING, chunks)
            print(generated_chunk)
        except Exception as e:
            print(e)
            winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
            exit(1)
        finally:
            save_response(generated_chunk)
        
        questions_split = [
            q.strip() for q in textwrap.dedent(generated_chunk).strip().split("--Question Starting--")
            if q.strip()
        ]
        
        if not TESTING and len(questions_split) != chunks:
            print(f"⚠️ GPT returned {len(questions_split)} questions instead of {chunks}. Skipping this chunk.")
            skipped_chunks.append(questions_split)
            continue

        cleaned_questions = [q.strip() for q in questions_split if q.strip()]
        all_questions.extend(cleaned_questions)

    # Shuffle for final verification
    random.shuffle(all_questions)

    # Save all raw questions to file BEFORE verification
    questions_output = "\n\n".join(all_questions).strip()
    try:
        save_to_docx(questions_output, QUESTIONS_PATH)
    except Exception as e:
        print(e)
        winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
        exit(1)

    verified_output = ""
    for i in range(0, len(all_questions), chunks):
        chunk = all_questions[i:i+chunks]
        if len(chunk) < chunks:
            print(f"⚠️ Skipping incomplete chunk at end of generation (less than 5 questions).")
            skipped_chunks.append(chunk)
            continue

        combined_text = "\n\n".join(chunk)

        # Verify the chunk via GPTtry:
        try:
            verification = call_gpt(verify_prompt(combined_text), TESTING, chunks)
        except Exception as e:
            print(e)
            winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
            exit(1)
        verified_output += verification + "\n\n"
    try:
        save_to_docx(verified_output.strip(), VERIFICATIONS_PATH)
    except Exception as e:
        print(e)
        winsound.PlaySound("WrongBuzzer.wav", winsound.SND_FILENAME)
        exit(1)
        
    if skipped_chunks:
        print("\n📭 Skipped Chunks (Incomplete):")
        for i, skipped in enumerate(skipped_chunks, 1):
            print(f"\n🔹 Skipped Chunk {i}:")
            print("\n\n".join(skipped))

        skipped_text = "\n\n".join(
            f"Skipped Chunk {i}:\n" + "\n\n".join(chunk)
            for i, chunk in enumerate(skipped_chunks, 1)
        )
        save_to_docx(skipped_text.strip(), SKIPPED_PATH)

    print("\n✅ Shuffled and verified mock paper saved.")
    winsound.PlaySound("CorrectHarp.wav", winsound.SND_FILENAME)

if __name__ == "__main__":
    GenerationMain(CHUNK_SIZE)